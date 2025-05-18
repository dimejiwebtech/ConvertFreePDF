import os
import time
import threading
from django.conf import settings
from django.http import JsonResponse, FileResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from pdf2image import convert_from_bytes
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from docx import Document
import uuid


# First Tool (PDF to DOC)

def pdf_to_doc(request):
    """Render the PDF to DOC conversion page."""
    return render(request, 'converters/pdf_to_doc.html')

@csrf_exempt
def convert_pdf_to_doc(request):
    """Convert PDF to DOCX."""
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]

        if not uploaded_file.name.endswith(".pdf"):
            return JsonResponse({"error": "Invalid file format"}, status=400)

        original_name = os.path.splitext(uploaded_file.name)[0]
        pdf_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.pdf")
        docx_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.docx")

        with open(pdf_path, "wb+") as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        try:
            pdf_text = extract_text_from_pdf(pdf_path)
            convert_to_docx(pdf_text, docx_path)
            os.remove(pdf_path)
            return JsonResponse({"success": True, "download_url": f"/tools/download/{original_name}.docx"})
        except Exception as e:
            return JsonResponse({"error": f"Conversion failed: {str(e)}"}, status=500)

    return JsonResponse({"error": "No file uploaded"}, status=400)

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    reader = PdfReader(pdf_path)
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def convert_to_docx(text, docx_path):
    """Convert extracted text into a DOCX file."""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(docx_path)

# Second Tool (DOC to PDF)

def doc_to_pdf(request):
    """Render the DOC to PDF conversion page."""
    return render(request, 'converters/doc_to_pdf.html')

@csrf_exempt
def convert_doc_to_pdf(request):
    """Convert DOCX to PDF."""
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]

        if uploaded_file.size > 50 * 1024 * 1024:
            return JsonResponse({"error": "File size exceeds 50MB limit"}, status=400)

        if not uploaded_file.name.endswith((".doc", ".docx")):
            return JsonResponse({"error": "Invalid file format"}, status=400)

        original_name = os.path.splitext(uploaded_file.name)[0]
        docx_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.docx")
        pdf_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.pdf")

        with open(docx_path, "wb+") as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        try:
            convert_to_pdf(docx_path, pdf_path)
            os.remove(docx_path)
            return JsonResponse({"success": True, "download_url": f"/tools/download/{original_name}.pdf"})
        except Exception as e:
            return JsonResponse({"error": f"Conversion failed: {str(e)}"}, status=500)

    return JsonResponse({"error": "No file uploaded"}, status=400)

def convert_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using pure Python."""
    doc = Document(docx_path)
    pdf_canvas = canvas.Canvas(pdf_path, pagesize=letter)
    y_position = letter[1] - 50

    for para in doc.paragraphs:
        if para.text.strip():
            pdf_canvas.drawString(100, y_position, para.text)
            y_position -= 20
            if y_position < 50:
                pdf_canvas.showPage()
                y_position = letter[1] - 50

    pdf_canvas.save()


def download_file(request, file_name):
    """Serve the converted file for download and delete it after a short delay."""
    file_path = os.path.join(settings.MEDIA_ROOT, file_name)

    if os.path.exists(file_path):
        response = FileResponse(open(file_path, "rb"), as_attachment=True)
        response["Content-Disposition"] = f'attachment; filename="{file_name}"'

        # Delete file after a short delay
        def delete_file():
            try:
                time.sleep(10)  # Wait 10 seconds before deleting
                os.remove(file_path)
                print(f"Deleted file: {file_path}")
            except Exception as e:
                print(f"Failed to delete {file_path}: {e}")

        threading.Thread(target=delete_file, daemon=True).start()
        return response

    return JsonResponse({"error": "File not found"}, status=404)

# Third Tool (PDF to JPG)

def pdf_to_jpg(request):
    return render(request, 'converters/pdf_to_jpg.html')

@csrf_exempt
def convert_pdf_to_jpg(request):
    """Handles PDF to JPG conversion using pdf2image without Poppler."""
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]

        # Validate file size (Max: 50MB)
        if uploaded_file.size > 50 * 1024 * 1024:
            return JsonResponse({"error": "File size exceeds the 50MB limit"}, status=400)

        # Validate file extension
        if not uploaded_file.name.endswith(".pdf"):
            return JsonResponse({"error": "Invalid file format. Please upload a PDF file."}, status=400)

        try:
            # Read PDF file into memory
            pdf_data = uploaded_file.read()
            images = convert_from_bytes(pdf_data)  # Convert PDF pages to images

            image_files = []
            original_name = os.path.splitext(uploaded_file.name)[0]

            # Save images as JPG
            for i, image in enumerate(images):
                jpg_filename = f"{original_name}_page_{i+1}.jpg"
                jpg_path = os.path.join(settings.MEDIA_ROOT, jpg_filename)
                image.save(jpg_path, "JPEG")
                image_files.append(jpg_filename)

            # Return download URLs
            download_urls = [f"/tools/download/{file}" for file in image_files]
            return JsonResponse({"success": True, "download_urls": download_urls})

        except Exception as e:
            return JsonResponse({"error": f"Conversion failed: {str(e)}"}, status=500)

    return JsonResponse({"error": "No file uploaded"}, status=400)

# Fourth Tool (JPG to PDF)

def jpg_to_pdf(request):
    return render(request, 'converters/jpg_to_pdf.html')

@csrf_exempt
def convert_jpg_to_pdf(request):
    """Convert JPG image(s) to a single PDF."""
    if request.method == "POST" and request.FILES.getlist("files"):
        uploaded_files = request.FILES.getlist("files")

        for file in uploaded_files:
            if not file.name.lower().endswith((".jpg", ".jpeg")):
                return JsonResponse({"error": "Invalid file format. Upload JPG files."}, status=400)

        original_name = os.path.splitext(uploaded_files[0].name)[0]
        pdf_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.pdf")

        try:
            images = [Image.open(file).convert("RGB") for file in uploaded_files]
            images[0].save(pdf_path, save_all=True, append_images=images[1:])
            return JsonResponse({"success": True, "download_url": f"/tools/download/{original_name}.pdf"})

        except Exception as e:
            return JsonResponse({"error": f"Conversion failed: {str(e)}"}, status=500)

    return JsonResponse({"error": "No files uploaded"}, status=400)

# Fifth Tool (PDF to PNG)

def pdf_to_png(request):
    return render(request, 'converters/pdf_to_png.html')

@csrf_exempt
def convert_pdf_to_png(request):
    """Convert PDF pages to PNG images."""
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]

        if not uploaded_file.name.endswith(".pdf"):
            return JsonResponse({"error": "Invalid file format. Upload a PDF."}, status=400)

        original_name = os.path.splitext(uploaded_file.name)[0]
        pdf_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.pdf")
        output_folder = os.path.join(settings.MEDIA_ROOT, f"{original_name}_png")

        with open(pdf_path, "wb+") as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        try:
            images = convert_from_path(pdf_path)
            os.makedirs(output_folder, exist_ok=True)
            image_paths = []

            for i, img in enumerate(images):
                img_path = os.path.join(output_folder, f"{original_name}_page_{i+1}.png")
                img.save(img_path, "PNG")
                image_paths.append(f"/tools/download/{os.path.basename(img_path)}")

            os.remove(pdf_path)
            return JsonResponse({"success": True, "download_urls": image_paths})

        except Exception as e:
            return JsonResponse({"error": f"Conversion failed: {str(e)}"}, status=500)

    return JsonResponse({"error": "No file uploaded"}, status=400)

# Sixth Tool (PDF Compressor)

def pdf_compressor(request):
    return render(request, 'converters/pdf_compressor.html')

@csrf_exempt
def compress_pdf(request):
    """Compress a PDF file by reducing its size."""
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]

        if not uploaded_file.name.endswith(".pdf"):
            return JsonResponse({"error": "Invalid file format. Upload a PDF."}, status=400)

        original_name = os.path.splitext(uploaded_file.name)[0]
        input_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.pdf")
        output_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}_compressed.pdf")

        with open(input_path, "wb+") as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        try:
            reader = PdfReader(input_path)
            writer = PdfWriter()

            for page in reader.pages:
                writer.add_page(page)

            with open(output_path, "wb") as output_pdf:
                writer.write(output_pdf)

            os.remove(input_path)
            return JsonResponse({"success": True, "download_url": f"/tools/download/{original_name}_compressed.pdf"})

        except Exception as e:
            return JsonResponse({"error": f"Compression failed: {str(e)}"}, status=500)

    return JsonResponse({"error": "No file uploaded"}, status=400)

# Seventh Tool (PDF Merger)

def merge_pdf_tool(request):
    """Render the merge PDF tool page."""
    return render(request, "converters/merge_pdf.html")

@csrf_exempt
def merge_pdfs(request):
    """Merge multiple PDF files into one."""
    if request.method == "POST" and request.FILES.getlist("files"):
        uploaded_files = request.FILES.getlist("files")

        for file in uploaded_files:
            if not file.name.endswith(".pdf"):
                return JsonResponse({"error": "Invalid file format. Upload PDF files only."}, status=400)

        original_name = "merged_document"
        merged_pdf_path = os.path.join(settings.MEDIA_ROOT, f"{original_name}.pdf")

        temp_paths = []

        try:
            merger = PdfMerger()

            # Save files temporarily with unique names to avoid name conflicts
            for file in uploaded_files:
                unique_filename = f"{uuid.uuid4().hex}.pdf"
                temp_path = os.path.join(settings.MEDIA_ROOT, unique_filename)

                with open(temp_path, "wb+") as destination:
                    for chunk in file.chunks():
                        destination.write(chunk)

                merger.append(temp_path)
                temp_paths.append(temp_path)

            # Write the merged PDF
            merger.write(merged_pdf_path)
            merger.close()

            # Clean up all temporary files
            for temp_file in temp_paths:
                os.remove(temp_file)

            return JsonResponse({"success": True, "download_url": f"/tools/download/{original_name}.pdf"})

        except Exception as e:
            # Clean up in case of error
            for temp_file in temp_paths:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            return JsonResponse({"error": f"Merge failed: {str(e)}"}, status=500)

    return JsonResponse({"error": "No files uploaded"}, status=400)

























